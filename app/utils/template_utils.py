from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io

from app.models.company import Company
from app.models.plan import PlanSlug
from app.services.exceptions import (
    FileStorageServiceException,
    FileNotFoundServiceException,
)


TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates" / "documents"
MAX_LOGO_WIDTH = 2.5
MAX_LOGO_HEIGHT = 1.5
MAX_HEADER_LOGO_WIDTH = 0.9
MAX_HEADER_LOGO_HEIGHT = 0.6


class TemplateLoader:
    @staticmethod
    def load_template(plan_slug: PlanSlug) -> Path:
        if not plan_slug:
            raise ValueError("plan_slug is required to load template")
        
        template_map = {
            PlanSlug.BASIC: "basic_manual_template.docx",
            PlanSlug.COMPREHENSIVE: "comprehensive_manual_template.docx",
            PlanSlug.INDUSTRY_SPECIFIC: "SJP-Template.docx",
        }
        
        template_filename = template_map.get(plan_slug)
        if not template_filename:
            raise ValueError(f"Unknown plan type: {plan_slug}")
        
        template_path = TEMPLATES_DIR / template_filename
        if not template_path.exists():
            raise FileNotFoundServiceException(f"Template not found: {template_path}")
        
        return template_path


def _replace_in_paragraph(paragraph, target: str, replacement: str) -> bool:
    if paragraph.runs:
        full_text = "".join(run.text for run in paragraph.runs)
        if target not in full_text:
            return False
        new_text = full_text.replace(target, replacement)
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""
        return True
    elif target in paragraph.text:
        paragraph.text = paragraph.text.replace(target, replacement)
        return True
    return False


def _replace_in_paragraphs(paragraphs, placeholder: str, value: str) -> None:
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            _replace_in_paragraph(paragraph, placeholder, value)


def _replace_in_tables(tables, placeholder: str, value: str) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, placeholder, value)


def replace_template_variables(doc: Document, replacements: dict[str, str]) -> Document:
    if not doc:
        raise ValueError("Document cannot be None")

    if not replacements:
        raise ValueError("Replacements dictionary cannot be empty")

    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        _replace_in_paragraphs(doc.paragraphs, placeholder, value)
        _replace_in_tables(doc.tables, placeholder, value)

        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                _replace_in_paragraphs(header_footer.paragraphs, placeholder, value)
                _replace_in_tables(header_footer.tables, placeholder, value)

    return doc


def resize_logo_image(
    logo_path: Path,
    max_width: float = MAX_LOGO_WIDTH,
    max_height: float = MAX_LOGO_HEIGHT,
) -> tuple[bytes, float, float]:
    if not logo_path:
        raise ValueError("logo_path is required")

    if not logo_path.exists():
        raise FileNotFoundServiceException(f"Logo file not found: {logo_path}")

    try:
        image = Image.open(logo_path)
    except Exception as e:
        raise FileStorageServiceException(f"Failed to open logo image: {str(e)}")

    original_width, original_height = image.size
    aspect_ratio = original_width / original_height

    if original_width > original_height:
        new_width = min(max_width, original_width / 96)
        new_height = new_width / aspect_ratio
        if new_height > max_height:
            new_height = max_height
            new_width = new_height * aspect_ratio
    else:
        new_height = min(max_height, original_height / 96)
        new_width = new_height * aspect_ratio
        if new_width > max_width:
            new_width = max_width
            new_height = new_width / aspect_ratio

    img_byte_arr = io.BytesIO()
    if logo_path.suffix.lower() == '.png':
        image.save(img_byte_arr, format='PNG')
    else:
        image.save(img_byte_arr, format='JPEG')
    img_byte_arr.seek(0)

    return img_byte_arr.getvalue(), new_width, new_height


def insert_company_logo(doc: Document, logo_path: Path, placeholder: str = "{{logo}}") -> Document:
    if not doc:
        raise ValueError("Document cannot be None")
    
    if not logo_path:
        raise ValueError("logo_path is required")
    
    if not placeholder:
        raise ValueError("placeholder is required")
    
    body_logo_bytes, body_width_inches, _ = resize_logo_image(logo_path)
    header_logo_bytes, header_width_inches, _ = resize_logo_image(
        logo_path,
        max_width=MAX_HEADER_LOGO_WIDTH,
        max_height=MAX_HEADER_LOGO_HEIGHT,
    )

    logo_inserted = False

    def _insert_in_paragraph(paragraph, is_header: bool) -> bool:
        if placeholder not in paragraph.text:
            return False
        _replace_in_paragraph(paragraph, placeholder, "")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        if is_header:
            run.add_picture(io.BytesIO(header_logo_bytes), width=Inches(header_width_inches))
        else:
            run.add_picture(io.BytesIO(body_logo_bytes), width=Inches(body_width_inches))
        return True

    for paragraph in doc.paragraphs:
        if _insert_in_paragraph(paragraph, is_header=False):
            logo_inserted = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _insert_in_paragraph(paragraph, is_header=False):
                        logo_inserted = True

    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for paragraph in header_footer.paragraphs:
                if _insert_in_paragraph(paragraph, is_header=True):
                    logo_inserted = True
            for table in header_footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if _insert_in_paragraph(paragraph, is_header=True):
                                logo_inserted = True

    if not logo_inserted:
        raise ValueError(f"Logo placeholder '{placeholder}' not found in document")
    
    return doc


def build_company_replacements(company: Company, order_id: int) -> dict[str, str]:
    if not company:
        raise ValueError("Company cannot be None")
    
    if not order_id:
        raise ValueError("order_id is required")
    
    company_name = company.name if company.name else "N/A"
    
    return {
        "company_name": company_name,
        "order_id": str(order_id),
        "generation_date": datetime.now().strftime("%B %d, %Y"),
        "year": str(datetime.now().year),
    }
