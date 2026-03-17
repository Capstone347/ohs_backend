from pathlib import Path
import subprocess
import shutil

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from pypdf import PdfReader, PdfWriter
import io

from app.models.plan import PlanSlug
from app.services.exceptions import FileStorageServiceException


def _find_libreoffice_binary() -> str:
    for candidate in ["libreoffice", "soffice", "/Applications/LibreOffice.app/Contents/MacOS/soffice"]:
        if shutil.which(candidate) or Path(candidate).exists():
            return candidate
    raise FileStorageServiceException(
        "LibreOffice is not installed. Install it with: brew install --cask libreoffice (macOS) "
        "or apt-get install libreoffice-writer (Linux)"
    )


def _find_cutoff_index(doc, plan_slug: PlanSlug) -> int | None:
    seen_toc = False
    heading_count = 0

    body = doc.element.body
    elements = list(body)

    for i, element in enumerate(elements):
        if element.tag.endswith('}p'):
            style_element = element.find(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle'
            )
            style_name = style_element.get(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', ''
            ) if style_element is not None else ''

            if style_name.startswith('toc') or style_name.startswith('TOC'):
                seen_toc = True
            elif seen_toc:
                return i

            if style_name == 'Heading1':
                heading_count += 1
                max_headings = 3 if plan_slug == PlanSlug.BASIC else 4
                if heading_count > max_headings:
                    return i

    return None


def create_limited_preview_docx(template_path: Path, plan_slug: PlanSlug) -> Path:
    if not template_path.exists():
        raise FileStorageServiceException(f"Template not found: {template_path}")

    doc = Document(str(template_path))

    cutoff = _find_cutoff_index(doc, plan_slug)

    if cutoff is not None:
        body = doc.element.body
        elements = list(body)
        for element in elements[cutoff:]:
            if element.tag.endswith('}sectPr'):
                continue
            body.remove(element)

    watermark_para = doc.add_paragraph()
    watermark_run = watermark_para.add_run("\n\n--- PREVIEW VERSION ---")
    watermark_run.font.size = Pt(16)
    watermark_run.font.bold = True
    watermark_run.font.color.rgb = RGBColor(200, 0, 0)
    watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    preview_para = doc.add_paragraph()
    preview_run = preview_para.add_run(
        "This is a limited preview. Purchase to access the complete manual with all sections, "
        "procedures, and templates."
    )
    preview_run.font.size = Pt(12)
    preview_run.italic = True
    preview_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    data_dir = Path(__file__).parent.parent.parent / "data" / "documents" / "previews"
    data_dir.mkdir(parents=True, exist_ok=True)
    preview_path = data_dir / f"preview_{template_path.name}"
    doc.save(str(preview_path))

    return preview_path


def docx_to_pdf(docx_path: Path) -> Path:
    if not docx_path.exists():
        raise FileStorageServiceException(f"DOCX file not found: {docx_path}")

    soffice = _find_libreoffice_binary()
    output_dir = str(docx_path.parent)

    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", output_dir, str(docx_path)],
        capture_output=True,
        text=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise FileStorageServiceException(f"LibreOffice PDF conversion failed: {result.stderr}")

    pdf_path = docx_path.with_suffix('.pdf')
    if not pdf_path.exists():
        raise FileStorageServiceException(f"PDF conversion produced no output: {pdf_path}")

    return pdf_path


def add_watermark_to_pdf(pdf_path: Path, watermark_text: str = "PREVIEW") -> Path:
    if not pdf_path.exists():
        raise FileStorageServiceException(f"PDF file not found: {pdf_path}")

    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()

    watermark_buffer = io.BytesIO()
    watermark_canvas = canvas.Canvas(watermark_buffer, pagesize=letter)
    width, height = letter

    watermark_canvas.saveState()
    watermark_canvas.setFont("Helvetica-Bold", 60)
    watermark_canvas.setFillColorRGB(0.9, 0.9, 0.9, alpha=0.3)
    watermark_canvas.translate(width / 2, height / 2)
    watermark_canvas.rotate(45)
    watermark_canvas.drawCentredString(0, 0, watermark_text)
    watermark_canvas.restoreState()
    watermark_canvas.save()

    watermark_buffer.seek(0)
    watermark_pdf = PdfReader(watermark_buffer)
    watermark_page = watermark_pdf.pages[0]

    for page in reader.pages:
        page.merge_page(watermark_page)
        writer.add_page(page)

    watermarked_path = pdf_path.parent / f"watermarked_{pdf_path.name}"
    with open(watermarked_path, 'wb') as output_file:
        writer.write(output_file)

    return watermarked_path


def create_secure_preview_pdf(
    template_path: Path,
    plan_slug: PlanSlug,
    replacements: dict[str, str],
    logo_path: Path | None = None,
) -> Path:
    from app.utils.template_utils import replace_template_variables, insert_company_logo

    limited_docx_path = create_limited_preview_docx(template_path, plan_slug)

    doc = Document(str(limited_docx_path))
    doc = replace_template_variables(doc, replacements)

    if logo_path:
        try:
            doc = insert_company_logo(doc, logo_path)
        except ValueError:
            pass
    else:
        doc = replace_template_variables(doc, {"logo": ""})

    doc.save(str(limited_docx_path))

    pdf_path = docx_to_pdf(limited_docx_path)

    watermarked_pdf_path = add_watermark_to_pdf(pdf_path, "PREVIEW - LIMITED CONTENT")

    limited_docx_path.unlink()
    pdf_path.unlink()

    return watermarked_pdf_path
