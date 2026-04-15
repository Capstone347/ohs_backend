import logging
import secrets
from datetime import UTC, datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models.document import Document as DocumentModel
from app.models.document import DocumentFormat
from app.models.sjp_content import SjpContentStatus
from app.repositories.document_repository import DocumentRepository
from app.repositories.sjp_content_repository import SjpContentRepository
from app.repositories.sjp_generation_job_repository import SjpGenerationJobRepository
from app.repositories.sjp_toc_entry_repository import SjpTocEntryRepository
from app.services.file_storage_service import FileStorageService
from app.services.jurisdiction_service import JurisdictionService

logger = logging.getLogger(__name__)


class SjpDocumentService:
    def __init__(
        self,
        sjp_job_repo: SjpGenerationJobRepository,
        sjp_toc_entry_repo: SjpTocEntryRepository,
        sjp_content_repo: SjpContentRepository,
        document_repo: DocumentRepository,
        file_storage_service: FileStorageService,
        jurisdiction_service: JurisdictionService,
    ):
        self.sjp_job_repo = sjp_job_repo
        self.sjp_toc_entry_repo = sjp_toc_entry_repo
        self.sjp_content_repo = sjp_content_repo
        self.document_repo = document_repo
        self.file_storage_service = file_storage_service
        self.jurisdiction_service = jurisdiction_service

    def generate_sjp_document(
        self,
        job_id: int,
        order_id: int,
        company_name: str,
        logo_path: Path | None = None,
    ) -> DocumentModel:
        job = self.sjp_job_repo.get_by_id(job_id)
        if not job:
            raise ValueError(f"SJP generation job {job_id} not found")

        toc_entries = self.sjp_toc_entry_repo.get_by_job_id(job_id)
        toc_ids = [entry.id for entry in toc_entries]
        contents = self.sjp_content_repo.get_by_toc_entry_ids(toc_ids) if toc_ids else []
        content_by_toc_id = {c.toc_entry_id: c for c in contents}

        completed_entries = [
            (entry, content_by_toc_id[entry.id])
            for entry in toc_entries
            if entry.id in content_by_toc_id
            and content_by_toc_id[entry.id].status == SjpContentStatus.COMPLETED.value
        ]

        if not completed_entries:
            raise ValueError(f"No completed SJP entries for job {job_id}")

        pack = self.jurisdiction_service.get_pack(job.province)

        doc = Document()
        self._set_default_font(doc)

        self._add_title_page(doc, company_name, job.province, pack.province_name)

        if logo_path and logo_path.exists():
            self._insert_logo(doc, logo_path)

        self._add_table_of_contents(doc, completed_entries)

        for entry, content in completed_entries:
            self._add_sjp_section(doc, entry, content, pack.province_name)

        self._add_disclaimer_page(doc, pack)

        output_path = self._save_document(doc, order_id)
        document_model = self._create_document_record(order_id, output_path)

        logger.info(
            "SJP document generated for order %d (job %d): %d SJPs",
            order_id, job_id, len(completed_entries),
        )

        return document_model

    def _set_default_font(self, doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_title_page(
        self, doc: Document, company_name: str, province_code: str, province_name: str
    ) -> None:
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("SAFE JOB PROCEDURES")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Industry-Specific Safety Manual")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(80, 80, 80)

        doc.add_paragraph()

        company = doc.add_paragraph()
        company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company.add_run(company_name)
        run.bold = True
        run.font.size = Pt(18)

        doc.add_paragraph()

        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info.add_run(f"Province: {province_name}")
        run.font.size = Pt(12)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)

        doc.add_page_break()

    def _insert_logo(self, doc: Document, logo_path: Path) -> None:
        try:
            doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = doc.paragraphs[0].add_run()
            run.add_picture(str(logo_path), width=Inches(2.5))
        except Exception:
            logger.warning("Failed to insert logo into SJP document")

    def _add_table_of_contents(self, doc: Document, entries: list) -> None:
        heading = doc.add_heading("Table of Contents", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        for entry, _ in entries:
            para = doc.add_paragraph()
            run = para.add_run(f"{entry.position}. {entry.title}")
            run.font.size = Pt(12)

        doc.add_page_break()

    def _add_sjp_section(self, doc: Document, entry, content, province_name: str) -> None:
        heading = doc.add_heading(f"SJP {entry.position}: {entry.title}", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        self._add_subsection(doc, "Task Description", content.task_description)
        self._add_list_subsection(doc, "Required PPE", content.required_ppe)
        self._add_numbered_subsection(doc, "Step-by-Step Instructions", content.step_by_step_instructions)
        self._add_list_subsection(doc, "Identified Hazards", content.identified_hazards)
        self._add_list_subsection(doc, "Control Measures", content.control_measures)
        self._add_list_subsection(doc, "Training Requirements", content.training_requirements)
        self._add_subsection(doc, "Emergency Procedures", content.emergency_procedures)

        if content.legislative_references:
            self._add_subsection(doc, "Legislative References", content.legislative_references, italic=True)

        doc.add_page_break()

    def _add_subsection(self, doc: Document, title: str, text: str, italic: bool = False) -> None:
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = RGBColor(51, 51, 51)

        para = doc.add_paragraph()
        run = para.add_run(text or "")
        run.italic = italic

    def _add_list_subsection(self, doc: Document, title: str, items: list) -> None:
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = RGBColor(51, 51, 51)

        if not items:
            doc.add_paragraph("N/A")
            return

        for item in items:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(str(item))

    def _add_numbered_subsection(self, doc: Document, title: str, items: list) -> None:
        heading = doc.add_heading(title, level=2)
        heading.runs[0].font.color.rgb = RGBColor(51, 51, 51)

        if not items:
            doc.add_paragraph("N/A")
            return

        for i, item in enumerate(items, start=1):
            para = doc.add_paragraph()
            para.add_run(f"{i}. {item}")

    def _add_disclaimer_page(self, doc: Document, pack) -> None:
        heading = doc.add_heading("Legal Disclaimer", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

        doc.add_paragraph()

        disclaimer_text = pack.prompt_preamble

        para = doc.add_paragraph()
        run = para.add_run(disclaimer_text)
        run.italic = True
        run.font.size = Pt(10)

        doc.add_paragraph()

        notice = doc.add_paragraph()
        run = notice.add_run(
            "This document was generated using artificial intelligence and has been reviewed "
            "by an administrator. While care has been taken to align the content with applicable "
            "occupational health and safety legislation, this document does not constitute legal "
            "advice. Employers should consult the current legislation and a qualified OHS "
            "professional to ensure full compliance."
        )
        run.italic = True
        run.font.size = Pt(10)

    def _save_document(self, doc: Document, order_id: int) -> Path:
        timestamp = int(datetime.now().timestamp())
        filename = f"sjp_order_{order_id}_{timestamp}.docx"
        output_path = self.file_storage_service.generated_documents_dir / filename
        doc.save(str(output_path))
        return output_path

    def _create_document_record(self, order_id: int, file_path: Path) -> DocumentModel:
        access_token = secrets.token_hex(32)
        token_expires_at = datetime.now(UTC) + timedelta(days=30)

        document = DocumentModel(
            order_id=order_id,
            file_path=str(file_path),
            file_format=DocumentFormat.DOCX.value,
            access_token=access_token,
            token_expires_at=token_expires_at,
            generated_at=datetime.now(UTC),
        )

        return self.document_repo.create(document)
