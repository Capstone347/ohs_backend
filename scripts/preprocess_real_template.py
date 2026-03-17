from pathlib import Path

from docx import Document


TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "documents"
SOURCE = TEMPLATES_DIR / "REAL_BASIC_MANUAL_TEMPLATE.docx"
OUTPUT = TEMPLATES_DIR / "basic_manual_template.docx"


def replace_across_runs(paragraph, target: str, replacement: str) -> bool:
    if paragraph.runs:
        full_text = "".join(run.text for run in paragraph.runs)
        if target not in full_text:
            return False
        new_text = full_text.replace(target, replacement)
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""
        return True
    elif target in paragraph.text:
        paragraph.text = paragraph.text.replace(target, replacement)
        return True
    return False


def process_paragraphs(paragraphs, target: str, replacement: str) -> int:
    count = 0
    for paragraph in paragraphs:
        if paragraph.runs:
            old_text = "".join(run.text for run in paragraph.runs)
        else:
            old_text = paragraph.text
        occurrences = old_text.count(target)
        if occurrences > 0:
            replace_across_runs(paragraph, target, replacement)
            count += occurrences
    return count


def process_tables(tables, target: str, replacement: str) -> int:
    count = 0
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                count += process_paragraphs(cell.paragraphs, target, replacement)
    return count


def process_headers_and_footers(doc, target: str, replacement: str) -> int:
    count = 0
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            count += process_paragraphs(header_footer.paragraphs, target, replacement)
            count += process_tables(header_footer.tables, target, replacement)
    return count


def preprocess():
    if not SOURCE.exists():
        raise FileNotFoundError(f"Source template not found: {SOURCE}")

    doc = Document(str(SOURCE))

    total = 0
    for area_fn in [
        lambda t, r: process_paragraphs(doc.paragraphs, t, r),
        lambda t, r: process_tables(doc.tables, t, r),
        lambda t, r: process_headers_and_footers(doc, t, r),
    ]:
        total += area_fn("COMPANY NAME", "{{company_name}}")

    print(f"Replaced 'COMPANY NAME' → '{{{{company_name}}}}': {total} occurrences")

    date_count = 0
    for section in doc.sections:
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = "".join(run.text for run in paragraph.runs)
                        if "March 9, 2026" in full_text:
                            replace_across_runs(paragraph, "March 9, 2026", "{{generation_date}}")
                            date_count += 1

    print(f"Replaced 'March 9, 2026' → '{{{{generation_date}}}}': {date_count} occurrences")

    logo_count = 0
    for section in doc.sections:
        for table in section.header.tables:
            if table.rows:
                cell = table.rows[0].cells[0]
                if cell.paragraphs:
                    paragraph = cell.paragraphs[0]
                    if not any("{{logo}}" in run.text for run in paragraph.runs):
                        if paragraph.runs:
                            paragraph.runs[0].text = "{{logo}}" + paragraph.runs[0].text
                        else:
                            paragraph.add_run("{{logo}}")
                        logo_count += 1

    print(f"Added '{{{{logo}}}}' placeholders: {logo_count}")

    doc.save(str(OUTPUT))
    print(f"\nSaved to: {OUTPUT}")

    verify_doc = Document(str(OUTPUT))
    placeholder_count = 0
    leftover_count = 0

    for paragraph in verify_doc.paragraphs:
        text = paragraph.text
        placeholder_count += text.count("{{company_name}}")
        leftover_count += text.count("COMPANY NAME")

    for table in verify_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    placeholder_count += text.count("{{company_name}}")
                    leftover_count += text.count("COMPANY NAME")

    for section in verify_doc.sections:
        for hf in [section.header, section.footer]:
            for paragraph in hf.paragraphs:
                text = paragraph.text
                placeholder_count += text.count("{{company_name}}")
                leftover_count += text.count("COMPANY NAME")
            for table in hf.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            placeholder_count += text.count("{{company_name}}")
                            leftover_count += text.count("COMPANY NAME")

    print(f"\nVerification:")
    print(f"  '{{{{company_name}}}}' count: {placeholder_count}")
    print(f"  'COMPANY NAME' leftover: {leftover_count}")

    assert placeholder_count > 200, f"Expected ~250 placeholders, got {placeholder_count}"
    assert leftover_count == 0, f"Found {leftover_count} leftover 'COMPANY NAME' occurrences"

    print("\nVerification passed!")


if __name__ == "__main__":
    preprocess()
