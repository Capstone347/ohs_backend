"""
Script to generate DOCX templates for the OHS Remote application.
Run this script to create the basic and comprehensive manual templates.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def create_basic_manual_template():
    doc = Document()
    
    # Title Section
    title = doc.add_heading('', level=0)
    title_run = title.add_run('{{logo}}')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    company_title = doc.add_heading('Occupational Health and Safety Manual', level=1)
    company_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_name = doc.add_paragraph()
    company_name_run = company_name.add_run('{{company_name}}')
    company_name_run.font.size = Pt(18)
    company_name_run.font.bold = True
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Document Information Section
    info_para = doc.add_paragraph()
    info_para.add_run('Order ID: ').bold = True
    info_para.add_run('{{order_id}}')
    
    date_para = doc.add_paragraph()
    date_para.add_run('Date Generated: ').bold = True
    date_para.add_run('{{generation_date}}')
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '2. Health and Safety Policy',
        '3. Responsibilities',
        '4. Hazard Identification and Risk Assessment',
        '5. Emergency Procedures',
        '6. Training Requirements',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        f'This Occupational Health and Safety (OHS) Manual has been developed for {{{{company_name}}}} '
        f'to establish a framework for maintaining a safe and healthy workplace. This manual outlines '
        f'our commitment to providing a safe working environment and defines the responsibilities of '
        f'all employees in maintaining workplace safety.'
    )
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'The purpose of this manual is to:'
    )
    purposes = [
        'Establish health and safety policies and procedures',
        'Define roles and responsibilities for workplace safety',
        'Provide guidance on hazard identification and control',
        'Ensure compliance with applicable occupational health and safety legislation',
    ]
    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 2: Health and Safety Policy
    doc.add_heading('2. Health and Safety Policy', level=1)
    doc.add_paragraph(
        f'{{{{company_name}}}} is committed to providing a safe and healthy working environment for all '
        f'employees, contractors, and visitors. We recognize that the health and safety of our workforce '
        f'is fundamental to our success and sustainability.'
    )
    
    doc.add_heading('2.1 Policy Statement', level=2)
    policy_box = doc.add_paragraph()
    policy_box.add_run(
        'Management is committed to:\n'
        '• Complying with all applicable health and safety legislation\n'
        '• Providing safe equipment and work procedures\n'
        '• Providing information, training, and supervision\n'
        '• Encouraging employee participation in health and safety matters\n'
        '• Continuously improving our health and safety performance'
    ).italic = True
    
    doc.add_page_break()
    
    # Section 3: Responsibilities
    doc.add_heading('3. Responsibilities', level=1)
    
    doc.add_heading('3.1 Management Responsibilities', level=2)
    doc.add_paragraph('Management is responsible for:')
    mgmt_resp = [
        'Establishing and maintaining the OHS management system',
        'Allocating resources for health and safety initiatives',
        'Setting health and safety objectives and targets',
        'Reviewing OHS performance regularly',
        'Leading by example in safety matters',
    ]
    for resp in mgmt_resp:
        doc.add_paragraph(resp, style='List Bullet')
    
    doc.add_heading('3.2 Supervisor Responsibilities', level=2)
    doc.add_paragraph('Supervisors are responsible for:')
    super_resp = [
        'Ensuring safe work procedures are followed',
        'Conducting workplace inspections',
        'Investigating incidents and near misses',
        'Providing on-the-job training and supervision',
        'Correcting unsafe conditions and behaviors',
    ]
    for resp in super_resp:
        doc.add_paragraph(resp, style='List Bullet')
    
    doc.add_heading('3.3 Employee Responsibilities', level=2)
    doc.add_paragraph('All employees are responsible for:')
    emp_resp = [
        'Working safely and following established procedures',
        'Using prescribed personal protective equipment (PPE)',
        'Reporting hazards, incidents, and near misses',
        'Participating in health and safety training',
        'Not engaging in activities that endanger themselves or others',
    ]
    for resp in emp_resp:
        doc.add_paragraph(resp, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 4: Hazard Identification
    doc.add_heading('4. Hazard Identification and Risk Assessment', level=1)
    doc.add_paragraph(
        f'{{{{company_name}}}} has established a systematic approach to identify workplace hazards '
        f'and assess associated risks. This process involves regular workplace inspections, '
        f'hazard reports from employees, and formal risk assessments.'
    )
    
    doc.add_heading('4.1 Hazard Identification Process', level=2)
    doc.add_paragraph('Hazards will be identified through:')
    hazard_methods = [
        'Regular workplace inspections',
        'Employee hazard reports',
        'Incident and near-miss investigations',
        'Job hazard analyses',
        'Pre-work safety assessments',
    ]
    for method in hazard_methods:
        doc.add_paragraph(method, style='List Bullet')
    
    doc.add_heading('4.2 Risk Assessment', level=2)
    doc.add_paragraph(
        'Once hazards are identified, risks will be assessed based on the likelihood of occurrence '
        'and potential severity of harm. Control measures will be implemented using the hierarchy of controls:'
    )
    controls = [
        'Elimination - Remove the hazard completely',
        'Substitution - Replace with a less hazardous alternative',
        'Engineering Controls - Isolate people from the hazard',
        'Administrative Controls - Change work procedures',
        'Personal Protective Equipment (PPE) - Protect the worker',
    ]
    for control in controls:
        doc.add_paragraph(control, style='List Number')
    
    doc.add_page_break()
    
    # Section 5: Emergency Procedures
    doc.add_heading('5. Emergency Procedures', level=1)
    
    doc.add_heading('5.1 Emergency Response Plan', level=2)
    doc.add_paragraph(
        f'{{{{company_name}}}} has developed emergency response procedures to ensure the safety '
        f'of all personnel in the event of an emergency. All employees must be familiar with '
        f'emergency procedures and evacuation routes.'
    )
    
    doc.add_heading('5.2 Emergency Contacts', level=2)
    
    # Create emergency contacts table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    contacts = [
        ('Emergency Services', '911'),
        ('Fire Department', 'Local Fire Department'),
        ('Police', 'Local Police Department'),
        ('Ambulance', 'Local Emergency Medical Services'),
        ('Poison Control', 'Local Poison Control Center'),
    ]
    
    for idx, (service, contact) in enumerate(contacts):
        row = table.rows[idx]
        row.cells[0].text = service
        row.cells[1].text = contact
    
    doc.add_heading('5.3 Evacuation Procedures', level=2)
    evac_steps = [
        'Upon hearing the alarm, immediately stop work',
        'Leave the building via the nearest safe exit',
        'Assist others who may need help',
        'Proceed to the designated assembly point',
        'Do not re-enter the building until authorized',
    ]
    for step in evac_steps:
        doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()
    
    # Section 6: Training
    doc.add_heading('6. Training Requirements', level=1)
    doc.add_paragraph(
        f'{{{{company_name}}}} is committed to providing comprehensive health and safety training '
        f'to all employees. Training ensures that workers have the knowledge and skills necessary '
        f'to perform their jobs safely.'
    )
    
    doc.add_heading('6.1 Mandatory Training', level=2)
    doc.add_paragraph('All employees must complete:')
    training_req = [
        'New employee orientation (including OHS overview)',
        'Job-specific safety training',
        'Emergency response procedures',
        'WHMIS (Workplace Hazardous Materials Information System)',
        'Personal protective equipment (PPE) training',
    ]
    for training in training_req:
        doc.add_paragraph(training, style='List Bullet')
    
    doc.add_heading('6.2 Training Records', level=2)
    doc.add_paragraph(
        'All safety training will be documented and records maintained. Training records will include:'
    )
    record_items = [
        'Employee name',
        'Training topic',
        'Date of training',
        'Trainer name',
        'Training duration',
    ]
    for item in record_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'{{{{company_name}}}} - OHS Manual | Generated: {{{{generation_date}}}} | © {{{{year}}}}'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def create_comprehensive_manual_template():
    doc = Document()
    
    # Title Section
    title = doc.add_heading('', level=0)
    title_run = title.add_run('{{logo}}')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    company_title = doc.add_heading('Comprehensive Occupational Health and Safety Manual', level=1)
    company_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    company_name = doc.add_paragraph()
    company_name_run = company_name.add_run('{{company_name}}')
    company_name_run.font.size = Pt(18)
    company_name_run.font.bold = True
    company_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Document Information Section
    info_para = doc.add_paragraph()
    info_para.add_run('Order ID: ').bold = True
    info_para.add_run('{{order_id}}')
    
    date_para = doc.add_paragraph()
    date_para.add_run('Date Generated: ').bold = True
    date_para.add_run('{{generation_date}}')
    
    version_para = doc.add_paragraph()
    version_para.add_run('Version: ').bold = True
    version_para.add_run('1.0')
    
    doc.add_page_break()
    
    # Extended Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction and Scope',
        '2. Health and Safety Policy Statement',
        '3. Organizational Structure and Responsibilities',
        '4. Hazard Identification and Risk Assessment',
        '5. Safe Work Procedures',
        '6. Emergency Response Plan',
        '7. Incident Investigation and Reporting',
        '8. Training and Competency',
        '9. Personal Protective Equipment (PPE)',
        '10. Workplace Inspections and Audits',
        '11. Contractor Management',
        '12. Document Control and Records',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Comprehensive sections (abbreviated for template)
    doc.add_heading('1. Introduction and Scope', level=1)
    doc.add_paragraph(
        f'This Comprehensive Occupational Health and Safety Manual provides {{{{company_name}}}} '
        f'with a complete framework for managing workplace health and safety. This manual applies to '
        f'all employees, contractors, visitors, and other persons at our workplace.'
    )
    
    doc.add_heading('1.1 Scope and Application', level=2)
    doc.add_paragraph(
        'This manual covers all aspects of health and safety management including but not limited to:'
    )
    scope_items = [
        'Workplace hazard identification and control',
        'Emergency preparedness and response',
        'Incident investigation and corrective actions',
        'Training and competency development',
        'Contractor and visitor safety management',
        'Compliance with applicable legislation and standards',
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('2. Health and Safety Policy Statement', level=1)
    doc.add_paragraph(
        f'{{{{company_name}}}} is committed to achieving the highest standards of health and safety '
        f'performance. We believe that all workplace incidents are preventable and that every employee '
        f'has the right to a safe and healthy work environment.'
    )
    
    doc.add_heading('2.1 Management Commitment', level=2)
    commitment_text = (
        'Senior management commits to:\n\n'
        '• Providing leadership and resources for health and safety initiatives\n'
        '• Setting measurable objectives and targets\n'
        '• Ensuring compliance with all applicable legislation\n'
        '• Fostering a culture of safety excellence\n'
        '• Continuously improving our OHS management system\n'
        '• Engaging and consulting with workers on safety matters\n'
        '• Holding all levels of management accountable for safety performance'
    )
    policy_para = doc.add_paragraph()
    policy_run = policy_para.add_run(commitment_text)
    policy_run.italic = True
    
    doc.add_page_break()
    
    doc.add_heading('3. Organizational Structure and Responsibilities', level=1)
    doc.add_paragraph(
        'Clear definition of roles and responsibilities is essential for effective health and safety management.'
    )
    
    # Add responsibility matrix table
    doc.add_heading('3.1 Responsibility Matrix', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'
    
    headers = ['Responsibility', 'Management', 'Supervisors', 'Employees']
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    responsibilities = [
        ('Policy Development', 'X', '', ''),
        ('Risk Assessment', 'X', 'X', ''),
        ('Training Delivery', '', 'X', ''),
        ('Hazard Reporting', '', 'X', 'X'),
        ('PPE Usage', '', '', 'X'),
    ]
    
    for idx, (resp, mgmt, super, emp) in enumerate(responsibilities, start=1):
        row = table.rows[idx]
        row.cells[0].text = resp
        row.cells[1].text = mgmt
        row.cells[2].text = super
        row.cells[3].text = emp
    
    doc.add_page_break()
    
    doc.add_heading('4. Hazard Identification and Risk Assessment', level=1)
    doc.add_paragraph(
        'A systematic approach to hazard identification and risk assessment is fundamental to '
        'preventing workplace incidents.'
    )
    
    doc.add_heading('4.1 Risk Assessment Methodology', level=2)
    doc.add_paragraph(
        'We use a risk matrix approach that considers both likelihood and severity:'
    )
    
    # Risk matrix table
    risk_table = doc.add_table(rows=5, cols=5)
    risk_table.style = 'Light Grid Accent 1'
    
    doc.add_paragraph()
    doc.add_paragraph('Risk Levels:')
    risk_levels = [
        'Low Risk - Acceptable with existing controls',
        'Medium Risk - Additional controls may be required',
        'High Risk - Immediate action required',
        'Extreme Risk - Activity must not proceed',
    ]
    for level in risk_levels:
        doc.add_paragraph(level, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('5. Safe Work Procedures', level=1)
    doc.add_paragraph(
        'Safe work procedures provide step-by-step instructions for performing tasks safely. '
        'All high-risk activities require documented safe work procedures.'
    )
    
    doc.add_heading('5.1 Procedure Development', level=2)
    doc.add_paragraph('Safe work procedures must include:')
    procedure_elements = [
        'Purpose and scope of the procedure',
        'Required personal protective equipment',
        'Step-by-step instructions',
        'Emergency procedures',
        'Responsibilities',
    ]
    for element in procedure_elements:
        doc.add_paragraph(element, style='List Number')
    
    doc.add_page_break()
    
    doc.add_heading('6. Emergency Response Plan', level=1)
    doc.add_paragraph(
        f'{{{{company_name}}}} maintains comprehensive emergency response procedures to protect '
        f'employees and minimize damage in emergency situations.'
    )
    
    doc.add_heading('6.1 Types of Emergencies', level=2)
    emergency_types = [
        'Fire and explosion',
        'Medical emergencies',
        'Hazardous material spills',
        'Natural disasters',
        'Workplace violence',
        'Utility failures',
    ]
    for emergency in emergency_types:
        doc.add_paragraph(emergency, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('7. Incident Investigation and Reporting', level=1)
    doc.add_paragraph(
        'All incidents, injuries, near misses, and hazards must be reported and investigated '
        'to prevent recurrence.'
    )
    
    doc.add_heading('7.1 Reporting Requirements', level=2)
    doc.add_paragraph('Immediate reporting is required for:')
    reporting_items = [
        'All injuries requiring first aid or medical attention',
        'Near miss incidents with potential for serious injury',
        'Property damage incidents',
        'Environmental releases',
        'Security incidents',
    ]
    for item in reporting_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('8. Training and Competency', level=1)
    doc.add_paragraph(
        'Training ensures that all personnel have the knowledge, skills, and competencies '
        'to perform their work safely.'
    )
    
    doc.add_heading('8.1 Training Matrix', level=2)
    training_table = doc.add_table(rows=6, cols=3)
    training_table.style = 'Light Grid Accent 1'
    
    training_headers = ['Training Course', 'Frequency', 'Target Audience']
    for idx, header in enumerate(training_headers):
        cell = training_table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    training_courses = [
        ('New Worker Orientation', 'Upon Hire', 'All Employees'),
        ('WHMIS', 'Every 3 Years', 'All Employees'),
        ('First Aid/CPR', 'Every 2 Years', 'Designated Personnel'),
        ('Fall Protection', 'Annual', 'At-Risk Workers'),
        ('Emergency Response', 'Annual', 'All Employees'),
    ]
    
    for idx, (course, freq, audience) in enumerate(training_courses, start=1):
        row = training_table.rows[idx]
        row.cells[0].text = course
        row.cells[1].text = freq
        row.cells[2].text = audience
    
    doc.add_page_break()
    
    doc.add_heading('9. Personal Protective Equipment (PPE)', level=1)
    doc.add_paragraph(
        'PPE is the last line of defense against workplace hazards. When other controls are '
        'not feasible, appropriate PPE must be used.'
    )
    
    doc.add_heading('9.1 PPE Assessment', level=2)
    doc.add_paragraph(
        'A PPE assessment will be conducted for all work areas to determine required equipment.'
    )
    
    doc.add_page_break()
    
    doc.add_heading('10. Workplace Inspections and Audits', level=1)
    doc.add_paragraph(
        'Regular inspections and audits verify compliance and identify opportunities for improvement.'
    )
    
    doc.add_heading('10.1 Inspection Schedule', level=2)
    inspection_items = [
        'Daily - Informal workplace observations',
        'Weekly - Formal workplace inspections',
        'Monthly - Equipment and machinery inspections',
        'Quarterly - Management system audits',
        'Annually - Comprehensive OHS program review',
    ]
    for item in inspection_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('11. Contractor Management', level=1)
    doc.add_paragraph(
        f'{{{{company_name}}}} ensures that contractors working at our facilities comply with our '
        f'health and safety standards.'
    )
    
    doc.add_heading('11.1 Contractor Requirements', level=2)
    contractor_reqs = [
        'Pre-qualification safety assessment',
        'Site-specific safety orientation',
        'Proof of insurance and certifications',
        'Compliance with safe work procedures',
        'Incident reporting obligations',
    ]
    for req in contractor_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_page_break()
    
    doc.add_heading('12. Document Control and Records', level=1)
    doc.add_paragraph(
        'This manual will be reviewed annually and updated as necessary to reflect changes in '
        'operations, legislation, or best practices.'
    )
    
    doc.add_heading('12.1 Record Retention', level=2)
    doc.add_paragraph('The following records will be maintained:')
    records = [
        'Training records - 3 years',
        'Incident investigations - 5 years',
        'Inspection reports - 2 years',
        'Risk assessments - Duration of activity plus 2 years',
        'Medical surveillance - Duration of employment plus 40 years',
    ]
    for record in records:
        doc.add_paragraph(record, style='List Bullet')
    
    doc.add_page_break()
    
    # Signature page
    doc.add_heading('Management Approval', level=1)
    doc.add_paragraph(
        'This Occupational Health and Safety Manual has been reviewed and approved by management.'
    )
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.rows[0].cells[0].text = 'Approved by:'
    sig_table.rows[0].cells[1].text = '_' * 40
    sig_table.rows[1].cells[0].text = 'Title:'
    sig_table.rows[1].cells[1].text = '_' * 40
    sig_table.rows[2].cells[0].text = 'Date:'
    sig_table.rows[2].cells[1].text = '_' * 40
    
    # Footer
    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'{{{{company_name}}}} - Comprehensive OHS Manual | Generated: {{{{generation_date}}}} | © {{{{year}}}} | Page '
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def main():
    templates_dir = Path(__file__).parent.parent / "templates" / "documents"
    templates_dir.mkdir(parents=True, exist_ok=True)
    
    print("Generating Basic Manual Template...")
    basic_doc = create_basic_manual_template()
    basic_path = templates_dir / "basic_manual_template.docx"
    basic_doc.save(str(basic_path))
    print(f"✓ Created: {basic_path}")
    
    print("\nGenerating Comprehensive Manual Template...")
    comprehensive_doc = create_comprehensive_manual_template()
    comprehensive_path = templates_dir / "comprehensive_manual_template.docx"
    comprehensive_doc.save(str(comprehensive_path))
    print(f"✓ Created: {comprehensive_path}")
    
    print("\n✓ All templates created successfully!")
    print(f"\nTemplates location: {templates_dir}")


if __name__ == "__main__":
    main()
